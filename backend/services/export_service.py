import os
import tempfile
import datetime
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document

def generate_pdf_report(title: str, content: str) -> str:
    """Generates a PDF report and returns the path to the temporary file."""
    if not title:
        title = "AI Generated Report"
        
    fd, path = tempfile.mkstemp(suffix=".pdf")
    os.close(fd)
    
    doc = SimpleDocTemplate(path, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = styles['Heading1']
    title_style.alignment = 1  # Center alignment
    
    date_style = ParagraphStyle(
        'DateStyle',
        parent=styles['Normal'],
        textColor='gray',
        spaceAfter=20,
        alignment=1  # Center alignment
    )
    
    content_style = styles['Normal']
    content_style.spaceAfter = 12
    
    elements = []
    
    # Add title
    elements.append(Paragraph(title, title_style))
    
    # Add date
    date_str = datetime.datetime.now().strftime("%B %d, %Y - %H:%M:%S")
    elements.append(Paragraph(f"Generated on {date_str}", date_style))
    elements.append(Spacer(1, 20))
    
    # Process content: handle basic line breaks
    paragraphs = content.split('\n')
    for p_text in paragraphs:
        if p_text.strip():
            # Replace basic markdown bold with HTML-like bold for reportlab
            text = p_text.replace('**', '<b>').replace('**', '</b>') # Simple replace might be buggy if odd number of **, but for basic needs it's okay. 
            # Better: count and replace properly or just remove ** for safety in simple export.
            # Let's just strip ** for safety if we don't have a markdown parser.
            text = p_text.replace('**', '')
            text = text.replace('*', '')
            elements.append(Paragraph(text.strip(), content_style))
        else:
            elements.append(Spacer(1, 12))
            
    doc.build(elements)
    return path

def generate_docx_report(title: str, content: str) -> str:
    """Generates a DOCX report and returns the path to the temporary file."""
    if not title:
        title = "AI Generated Report"
        
    doc = Document()
    
    # Title
    doc.add_heading(title, 0)
    
    # Date
    date_str = datetime.datetime.now().strftime("%B %d, %Y - %H:%M:%S")
    doc.add_paragraph(f"Generated on {date_str}", style='Subtitle')
    
    # Content
    paragraphs = content.split('\n')
    for p_text in paragraphs:
        text = p_text.strip()
        if text:
            # Strip simple markdown for plain docx
            clean_text = text.replace('**', '').replace('*', '')
            doc.add_paragraph(clean_text)
            
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    
    doc.save(path)
    return path
